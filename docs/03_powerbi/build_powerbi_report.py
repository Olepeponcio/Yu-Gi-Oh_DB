from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

BASE_DIR = Path(__file__).resolve().parent
OUT_PATH = BASE_DIR / "informe_analisis_powerbi_yugioh.docx"
MODEL_SVG = BASE_DIR / "modelo_relacional.svg"
MODEL_IMG = BASE_DIR / "modelo_relacional_powerbi.png"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"
MUTED = "666666"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def style_table(table, widths: list[int] | None = None, header: bool = True) -> None:
    set_table_width(table)
    table.autofit = False
    if table.rows:
        mark_header_row(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if widths:
                cell.width = Inches(widths[cell_idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            set_cell_margins(cell)
            if header and row_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Actualizar campo en Word"
    fld_sep.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_after = Pt(6)


def add_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Informe Power BI - Yu-Gi-Oh! | Pagina ")
    add_field(footer, "PAGE")


def add_metadata_table(doc: Document) -> None:
    rows = [
        ("Proyecto", "Proyecto SQL DB Yu-Gi-Oh"),
        ("Documento", "Informe de diseño y documentacion del analisis en Power BI"),
        ("Autor", "Pepin"),
        ("Fecha", "Junio 2026"),
        ("Version", "0.1 - estructura inicial"),
        (
            "Estado",
            "Documento vivo: se completara con capturas, medidas DAX y hallazgos.",
        ),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for idx, (label, value) in enumerate(rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
    style_table(table, widths=[2500, 6860], header=False)
    for row in table.rows:
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def add_status(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = text
    style_table(table, widths=[9360], header=False)
    set_cell_shading(table.cell(0, 0), "F8FAFC")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_kv_table(
    doc: Document, rows: list[tuple[str, str]], widths=(2600, 6760)
) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Elemento"
    table.cell(0, 1).text = "Detalle"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    style_table(table, widths=list(widths), header=True)


def add_matrix(
    doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, widths=widths, header=True)


def draw_model_image() -> None:
    img = Image.new("RGB", (1800, 1150), "#F6F7FB")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arial.ttf", 34)
        box_font = ImageFont.truetype("arial.ttf", 20)
        small_font = ImageFont.truetype("arial.ttf", 16)
        tiny_font = ImageFont.truetype("arial.ttf", 13)
    except OSError:
        title_font = ImageFont.load_default()
        box_font = ImageFont.load_default()
        small_font = ImageFont.load_default()
        tiny_font = ImageFont.load_default()

    draw.text(
        (70, 48),
        "Modelo relacional de vistas para Power BI",
        fill="#172033",
        font=title_font,
    )
    draw.text(
        (70, 92),
        "Resultado simplificado: dimensiones y hechos base cargados como vistas para analisis.",
        fill="#5A6475",
        font=small_font,
    )

    boxes = {
        "cards": (
            70,
            165,
            430,
            335,
            "vw_dim_cards_descriptive\nGrano: 1 carta\nClave: card_id",
            "dim",
        ),
        "sets": (
            70,
            390,
            430,
            560,
            "vw_dim_sets_descriptive\nGrano: 1 set\nClave: set_id",
            "dim",
        ),
        "rarities": (
            70,
            615,
            430,
            785,
            "vw_dim_rarities_descriptive\nGrano: 1 rareza\nClave: rarity_id",
            "dim",
        ),
        "prices": (
            775,
            160,
            1115,
            340,
            "vw_fact_card_prices_descriptive\nGrano: carta + marketplace\nMedida: price",
            "fact",
        ),
        "appearances": (
            775,
            420,
            1115,
            600,
            "vw_fact_card_set_appearances\nGrano: carta + set + rareza\nMedidas: set_price, appearance_count",
            "bridge",
        ),
        "variation": (
            775,
            650,
            1115,
            840,
            "vw_fact_card_price_variation_predictive\nGrano: carta + marketplace + snapshot\nMedidas: price, price_change",
            "fact",
        ),
        "markets": (
            1390,
            165,
            1730,
            335,
            "vw_dim_marketplaces_descriptive\nGrano: 1 marketplace\nClave: marketplace",
            "dim",
        ),
        "currencies": (
            1390,
            390,
            1730,
            560,
            "vw_dim_currencies_descriptive\nGrano: 1 moneda\nClave: currency",
            "dim",
        ),
        "snapshots": (
            1390,
            615,
            1730,
            785,
            "vw_dim_snapshots_descriptive\nGrano: 1 snapshot\nClave: snapshot_at",
            "date",
        ),
    }

    def port(key, side, offset=0):
        x1, y1, x2, y2, _, _ = boxes[key]
        if side == "top":
            return ((x1 + x2) // 2 + offset, y1)
        if side == "bottom":
            return ((x1 + x2) // 2 + offset, y2)
        if side == "left":
            return (x1, (y1 + y2) // 2 + offset)
        return (x2, (y1 + y2) // 2 + offset)

    def connector(start, end, color="#536073"):
        draw.line([start, end], fill=color, width=4)
        x1, y1 = start
        x2, y2 = end
        dx = x2 - x1
        dy = y2 - y1
        if abs(dx) >= abs(dy):
            arrow = (
                [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)]
                if dx > 0
                else [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
            )
        else:
            arrow = (
                [(x2, y2), (x2 - 8, y2 - 14), (x2 + 8, y2 - 14)]
                if dy > 0
                else [(x2, y2), (x2 - 8, y2 + 14), (x2 + 8, y2 + 14)]
            )
        draw.polygon(arrow, fill=color)

    def cardinality(one_xy, many_xy):
        draw.text(one_xy, "1", fill="#050B18", font=box_font)
        draw.text(many_xy, "*", fill="#050B18", font=box_font)

    connector(port("cards", "right", -35), port("prices", "left", -35), "#2368A2")
    cardinality((455, 210), (745, 220))
    connector(port("cards", "right", 30), port("appearances", "left", -45), "#2368A2")
    cardinality((455, 275), (745, 450))
    connector(port("cards", "right", 70), port("variation", "left", 35), "#2368A2")
    cardinality((455, 315), (745, 760))
    connector(port("sets", "right", 0), port("appearances", "left", 0), "#536073")
    cardinality((455, 480), (745, 510))
    connector(port("rarities", "right", 0), port("appearances", "left", 45), "#536073")
    cardinality((455, 700), (745, 565))
    connector(port("markets", "left", -35), port("prices", "right", -35), "#2368A2")
    cardinality((1348, 200), (1120, 220))
    connector(port("currencies", "left", -20), port("prices", "right", 35), "#2368A2")
    cardinality((1348, 430), (1120, 285))
    connector(port("markets", "left", 35), port("variation", "right", -40), "#2368A2")
    cardinality((1348, 270), (1120, 745))
    connector(port("currencies", "left", 30), port("variation", "right", 10), "#2368A2")
    cardinality((1348, 480), (1120, 795))
    connector(port("snapshots", "left", 0), port("variation", "right", 55), "#6B4AB6")
    cardinality((1348, 705), (1120, 840))

    def box(key):
        x1, y1, x2, y2, text, kind = boxes[key]
        colors = {
            "dim": ("#EAF3FF", "#2368A2"),
            "fact": ("#EAF8EE", "#2D7C43"),
            "bridge": ("#FFF6E5", "#A66A00"),
            "date": ("#F1ECFF", "#6B4AB6"),
        }
        fill, outline = colors[kind]
        draw.rounded_rectangle(
            (x1, y1, x2, y2), radius=14, fill=fill, outline=outline, width=3
        )
        lines = text.split("\n")
        y = y1 + 28
        for i, line in enumerate(lines):
            font = box_font if i == 0 else small_font
            draw.text((x1 + 20, y), line, fill="#182235", font=font)
            y += 36 if i == 0 else 30

    for key in boxes:
        box(key)

    panel_x, panel_y, panel_w, panel_h = 70, 880, 760, 220
    draw.rounded_rectangle(
        (panel_x, panel_y, panel_x + panel_w, panel_y + panel_h),
        radius=12,
        fill="#FFFFFF",
        outline="#C8CEDA",
        width=2,
    )
    draw.text(
        (panel_x + 20, panel_y + 18),
        "Tabla de relaciones",
        fill="#182235",
        font=box_font,
    )
    rows = [
        ("vw_dim_cards_descriptive", "precios, apariciones, variacion", "1 : N"),
        ("vw_dim_sets_descriptive", "vw_fact_card_set_appearances", "1 : N"),
        ("vw_dim_rarities_descriptive", "vw_fact_card_set_appearances", "1 : N"),
        ("marketplace / currency / snapshot", "hechos de precio", "1 : N"),
    ]
    y = panel_y + 58
    for i, row in enumerate(rows):
        fill = "#F0F3F8" if i % 2 else "#FFFFFF"
        draw.rectangle((panel_x + 20, y - 6, panel_x + panel_w - 20, y + 26), fill=fill)
        draw.text((panel_x + 35, y), row[0], fill="#253044", font=tiny_font)
        draw.text((panel_x + 300, y), row[1], fill="#253044", font=tiny_font)
        draw.text((panel_x + 640, y), row[2], fill="#253044", font=tiny_font)
        y += 36

    legend_x, legend_y = 1010, 910
    draw.rounded_rectangle(
        (legend_x, legend_y, 1730, 1060),
        radius=12,
        fill="#FFFFFF",
        outline="#C8CEDA",
        width=2,
    )
    draw.text(
        (legend_x + 25, legend_y + 30), "Catalogo visual", fill="#182235", font=box_font
    )
    for idx, (label, color) in enumerate(
        [
            ("Dimension descriptiva", "#2368A2"),
            ("Hecho de precios", "#2D7C43"),
            ("Hecho puente carta-set-rareza", "#A66A00"),
            ("Dimension temporal snapshot", "#6B4AB6"),
        ]
    ):
        x = legend_x + 25 + (idx % 2) * 300
        y = legend_y + 70 + (idx // 2) * 42
        draw.rounded_rectangle((x, y, x + 24, y + 16), radius=4, fill=color)
        draw.text((x + 35, y - 2), label, fill="#4D586B", font=small_font)

    img.save(MODEL_IMG)


def build_doc() -> None:
    draw_model_image()
    doc = Document()
    configure_styles(doc)
    add_footer(doc.sections[0])
    set_update_fields(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Analisis del Mercado de Cartas Yu-Gi-Oh!")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Informe de diseño y documentación del análisis en Power BI")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()
    add_metadata_table(doc)
    doc.add_paragraph()
    add_status(
        doc,
        "Proposito del documento: acompanar al dashboard de Power BI. El panel responde que ocurre; "
        "este informe documenta por que se analiza, como se ha construido el modelo y que conclusiones "
        "se podran defender cuando existan visuales y medidas validadas.",
    )

    doc.add_page_break()
    doc.add_heading("Indice", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_paragraph(
        "Nota: en Word, actualizar campos para regenerar el indice cuando el documento evolucione."
    )

    doc.add_page_break()
    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Este informe documenta el diseño del analisis en Power BI para un proyecto de datos de cartas "
        "Yu-Gi-Oh. El origen operativo es una extraccion desde la API publica YGOPRODeck, normalizada "
        "mediante un ETL en Python y cargada en MySQL bajo el schema yugioh_db."
    )
    doc.add_paragraph(
        "La fase actual deja preparado el marco de consumo analitico: dimensiones y hechos base en formato `vw_`. "
        "Power BI consumira esas vistas sin modificar las tablas madre."
    )
    add_status(
        doc,
        "Estado actual: el informe ya contiene estructura, arquitectura, reglas de modelado, vistas de consumo "
        "y criterios de lectura. Quedan pendientes las capturas del dashboard, medidas DAX definitivas, resultados "
        "cuantitativos y conclusiones basadas en visuales.",
    )

    doc.add_heading("2. Objetivos del proyecto", level=1)
    add_bullets(
        doc,
        [
            "Analizar el catalogo de cartas Yu-Gi-Oh disponible desde YGOPRODeck.",
            "Construir un modelo Power BI trazable desde MySQL y no desde reglas improvisadas en visuales.",
            "Describir precios actuales por carta, marketplace y moneda.",
            "Analizar sets, rarezas, apariciones y precios desde hechos base.",
            "Preparar una base para recomendaciones prescriptivas revisables, no basadas en rankings aislados.",
        ],
    )

    doc.add_heading("3. Descripcion de los datos", level=1)
    add_kv_table(
        doc,
        [
            ("Origen externo", "API publica YGOPRODeck."),
            ("Origen interno para Power BI", "MySQL schema yugioh_db."),
            (
                "Extraccion",
                "src/api/ygoprodeck_client.py guarda data/raw/cardinfo_latest.json.",
            ),
            ("Carga", "ETL Python hacia tablas madre definidas en sql/schema.sql."),
            (
                "Numero de registros",
                "Pendiente: depende de la ejecucion local del ETL y de la fecha de extraccion.",
            ),
            ("Numero de tablas madre", "10 tablas principales."),
            (
                "Historico",
                "card_price_history inserta snapshots de precios en cada carga real del ETL.",
            ),
        ],
    )

    doc.add_heading("3.1 Tablas madre", level=2)
    add_matrix(
        doc,
        ["Tabla", "Funcion", "Grano principal"],
        [
            ["cards", "Catalogo base de cartas.", "1 carta"],
            ["sets", "Catalogo de sets.", "1 set"],
            [
                "rarities",
                "Catalogo tecnico de rarezas por codigo.",
                "1 set_code + rareza + codigo",
            ],
            [
                "card_sets",
                "Apariciones de cartas en sets.",
                "1 carta + 1 set/codigo + 1 rareza",
            ],
            ["card_images", "Imagenes asociadas.", "1 imagen"],
            ["card_prices", "Precios actuales por marketplace.", "1 carta"],
            [
                "card_price_history",
                "Snapshots historicos de precios.",
                "1 carta + 1 snapshot",
            ],
            ["card_banlist", "Estado de banlist.", "1 carta"],
            ["card_typelines", "Typelines por carta.", "1 carta + 1 typeline"],
            [
                "card_linkmarkers",
                "Marcadores Link por carta.",
                "1 carta + 1 linkmarker",
            ],
        ],
        [1900, 4300, 3160],
    )

    doc.add_heading("4. Preparacion de datos", level=1)
    doc.add_paragraph(
        "La preparacion se realiza antes de Power BI. El programa Python extrae el JSON, conserva una copia raw, "
        "normaliza dominios y carga MySQL. Power BI debe consumir consultas documentadas, no reconstruir la logica "
        "de negocio desde tablas crudas sin control."
    )
    add_numbered(
        doc,
        [
            "Descargar JSON o leer un JSON local.",
            "Normalizar cartas, sets, rarezas, apariciones, precios, imagenes, banlist y relaciones.",
            "Validar claves y campos requeridos.",
            "Insertar o actualizar tablas madre.",
            "Registrar snapshot de precios en card_price_history.",
            "Publicar vistas SQL de consumo para Power BI.",
        ],
    )

    doc.add_heading("4.1 Reglas criticas de preparacion", level=2)
    add_bullets(
        doc,
        [
            "No mezclar monedas sin segmentacion o conversion explicita.",
            "cardmarket_price llega en EUR; tcgplayer_price, ebay_price, amazon_price y coolstuffinc_price llegan en USD.",
            "set_price pertenece a card_sets; no es un precio propio de la rareza.",
            "Los hechos base cargados en Power BI son precios actuales, apariciones carta-set-rareza y variacion historica.",
            "Los rankings, outliers y resumenes se calculan como medidas o filtros desde hechos base.",
        ],
    )

    doc.add_heading("4.2 Procesos aplicados para el informe Power BI", level=2)
    add_matrix(
        doc,
        ["Proceso", "Aplicacion", "Criterio vigente"],
        [
            [
                "Carga y transformacion ETL",
                "Datos extraidos, transformados y cargados en MySQL desde el flujo Python.",
                "Power BI consume las vistas SQL publicadas; no modifica tablas madre.",
            ],
            [
                "Nulos en atributos numericos de carta",
                "`atk`, `def` y `link_value` se interpretan como 0 cuando no aplican; `scale` se interpreta como -1.",
                "Evita mezclar nulos tecnicos con valores analizables en segmentaciones y tarjetas.",
            ],
            [
                "Variacion porcentual de precio",
                "`vw_fact_card_price_variation_predictive.price_change_pct` conserva `NULL` cuando no hay base comparable.",
                "Las metricas y visuales aplican contexto de filtro para excluir valores vacios cuando el calculo lo requiera.",
            ],
            [
                "Filtrado aplicado en Power BI",
                "Las filas con `price_change_pct` vacio se excluyen en pasos aplicados o filtros del visual cuando se analizan variaciones.",
                "La vista SQL conserva el dato original; el filtro pertenece al contexto analitico.",
            ],
        ],
        [2500, 4300, 2560],
    )

    doc.add_heading("5. Arquitectura y modelo de datos", level=1)
    doc.add_paragraph(
        "El modelo recomendado para Power BI es una constelacion simplificada de hechos. Las dimensiones filtran "
        "hechos con direccion unica 1 -> *. No se recomienda relacionar hechos entre si, porque aumenta el riesgo "
        "de filtros ambiguos y dobles conteos."
    )
    picture = doc.add_picture(str(MODEL_IMG), width=Cm(16.0))
    picture._inline.docPr.set("title", "Modelo Power BI recomendado")
    picture._inline.docPr.set(
        "descr",
        "Diagrama de constelacion de hechos con dimensiones de cartas, marketplaces, monedas, snapshots y rarezas.",
    )
    caption = doc.add_paragraph(
        "Figura 1. Modelo Power BI recomendado a partir de las vistas de consumo."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("5.1 Relaciones recomendadas", level=2)
    add_bullets(
        doc,
        [
            "vw_dim_cards_descriptive 1 -> * vw_fact_card_prices_descriptive.",
            "vw_dim_cards_descriptive 1 -> * vw_fact_card_set_appearances.",
            "vw_dim_cards_descriptive 1 -> * vw_fact_card_price_variation_predictive.",
            "vw_dim_sets_descriptive 1 -> * vw_fact_card_set_appearances.",
            "vw_dim_rarities_descriptive 1 -> * vw_fact_card_set_appearances.",
            "vw_dim_marketplaces_descriptive 1 -> * hechos de precios y variacion.",
            "vw_dim_currencies_descriptive 1 -> * hechos de precios y variacion.",
            "vw_dim_snapshots_descriptive 1 -> * vw_fact_card_price_variation_predictive.",
        ],
    )

    doc.add_heading("6. Vistas de consumo para Power BI", level=1)
    add_matrix(
        doc,
        ["Vista", "Tipo", "Uso recomendado"],
        [
            ["vw_dim_cards_descriptive", "Dimension", "Catalogo base de cartas."],
            ["vw_dim_sets_descriptive", "Dimension", "Catalogo de sets."],
            [
                "vw_dim_rarities_descriptive",
                "Dimension",
                "Catalogo tecnico de rarezas.",
            ],
            [
                "vw_dim_marketplaces_descriptive",
                "Dimension",
                "Segmentacion de fuentes de precio.",
            ],
            ["vw_dim_currencies_descriptive", "Dimension", "Segmentacion por moneda."],
            [
                "vw_dim_snapshots_descriptive",
                "Dimension",
                "Calendario real de historico disponible.",
            ],
            [
                "vw_fact_card_prices_descriptive",
                "Hecho",
                "Precios actuales en formato largo.",
            ],
            [
                "vw_fact_card_set_appearances",
                "Hecho puente",
                "Apariciones carta-set-rareza y precio de set.",
            ],
            [
                "vw_fact_card_price_variation_predictive",
                "Hecho historico",
                "Variacion temporal por carta, marketplace y moneda.",
            ],
        ],
        [3950, 1950, 3460],
    )

    doc.add_heading("7. Metodologia analitica", level=1)
    add_numbered(
        doc,
        [
            "Analisis descriptivo: entender que existe, como se distribuye y que cobertura tienen los datos.",
            "Analisis diagnostico: explicar por que ciertos precios, rarezas o cartas destacan.",
            "Analisis predictivo: estudiar variacion temporal solo si hay snapshots suficientes.",
            "Analisis prescriptivo: convertir criterios validados en decisiones revisables.",
        ],
    )

    doc.add_heading("8. Analisis descriptivo", level=1)
    doc.add_paragraph(
        "Esta seccion se completara siguiendo el orden de las paginas del dashboard. Cada apartado debe mantener "
        "el patron: visual, que muestra, interpretacion y conclusion."
    )
    add_matrix(
        doc,
        ["Apartado", "Vista base", "Estado"],
        [
            ["Catalogo de cartas disponibles", "vw_dim_cards_descriptive", "Preparado"],
            [
                "Distribucion por tipo de carta",
                "vw_dim_cards_descriptive",
                "Preparado",
            ],
            [
                "Apariciones por set y rareza",
                "vw_fact_card_set_appearances",
                "Preparado",
            ],
            [
                "Distribucion de precios actuales",
                "vw_fact_card_prices_descriptive",
                "Preparado",
            ],
        ],
        [2800, 4200, 2360],
    )

    doc.add_heading("9. Analisis diagnostico", level=1)
    add_matrix(
        doc,
        ["Pregunta", "Consulta", "Criterio"],
        [
            [
                "Que cartas aparecen en mas sets",
                "vw_fact_card_set_appearances",
                "Interpretar como disponibilidad o reimpresion.",
            ],
            [
                "Que relacion existe entre rareza y precio",
                "vw_fact_card_set_appearances",
                "Agregar con grano controlado; no atribuir set_price a rarities.",
            ],
            [
                "Que precios requieren revision",
                "vw_fact_card_prices_descriptive",
                "Aplicar medidas/filtros de revision sobre precio y moneda.",
            ],
        ],
        [2600, 3800, 2960],
    )

    doc.add_heading("10. Analisis predictivo", level=1)
    doc.add_paragraph(
        "El analisis predictivo queda condicionado por el numero de snapshots reales en card_price_history. "
        "Sin historico suficiente no debe hablarse de tendencia; solo de disponibilidad de datos."
    )
    add_bullets(
        doc,
        [
            "vw_dim_snapshots_descriptive: permite validar fechas de snapshot disponibles.",
            "vw_fact_card_price_variation_predictive: compara precio actual contra snapshot anterior por carta, marketplace y moneda.",
            "Requiere conservar moneda y marketplace durante todo el analisis.",
        ],
    )

    doc.add_heading("11. Analisis prescriptivo", level=1)
    doc.add_paragraph(
        "La parte prescriptiva debe esperar a que descriptivo, diagnostico y predictivo esten validados. "
        "No se deben convertir rankings aislados en recomendaciones."
    )
    add_matrix(
        doc,
        ["Decision futura", "Criterio de avance", "Estado"],
        [
            [
                "Seleccionar carta principal potencial",
                "Combinar valor, presencia y legalidad.",
                "Pendiente",
            ],
            [
                "Seleccionar carta complementaria",
                "Precio moderado y coherencia tematica.",
                "Pendiente",
            ],
            [
                "Marcar carta para revision",
                "Outlier, moneda mezclada o dato incompleto.",
                "Pendiente",
            ],
        ],
        [3000, 4300, 2060],
    )

    doc.add_heading("12. Paginas previstas del dashboard", level=1)
    add_matrix(
        doc,
        ["Pagina", "Pregunta base", "Estado"],
        [
            [
                "Vista general",
                "Que contiene el catalogo y que fuentes de precio existen",
                "Pendiente",
            ],
            [
                "Precios actuales",
                "Como se distribuyen los precios por marketplace y moneda",
                "Pendiente",
            ],
            [
                "Sets y rarezas",
                "Que sets, apariciones y rarezas explican diferencias",
                "Pendiente",
            ],
            [
                "Revision de precios",
                "Que precios requieren revision antes de interpretarse",
                "Pendiente",
            ],
            [
                "Historico de precios",
                "Existe variacion temporal comparable",
                "Pendiente",
            ],
            [
                "Revision prescriptiva",
                "Que cartas merecen seguimiento o descarte",
                "Pendiente",
            ],
        ],
        [2300, 5000, 2060],
    )

    doc.add_heading("13. Conclusiones iniciales", level=1)
    add_bullets(
        doc,
        [
            "El proyecto ya separa responsabilidades: Python carga, MySQL conserva y Power BI consume vistas documentadas.",
            "El modelo BI debe tratarse como constelacion de hechos con dimensiones compartidas.",
            "El precio debe analizarse siempre con moneda declarada; mezclar EUR y USD invalida comparaciones directas.",
            "Los outliers son una lista de revision, no una conclusion automatica.",
            "La fase prescriptiva requiere reglas validadas antes de emitir recomendaciones.",
        ],
    )

    doc.add_heading("14. Limitaciones", level=1)
    add_bullets(
        doc,
        [
            "No se han incluido todavia resultados cuantitativos de una ejecucion local concreta.",
            "No hay capturas del dashboard ni medidas DAX definitivas en esta version.",
            "Los precios proceden de campos de marketplaces y pueden cambiar con cada actualizacion.",
            "No se consideran costes de envio, estado fisico de la carta, ventas privadas ni liquidez real.",
            "El historico depende de ejecuciones reales del ETL; sin suficientes snapshots no hay tendencia robusta.",
        ],
    )

    doc.add_heading("15. Trabajo futuro", level=1)
    add_bullets(
        doc,
        [
            "Completar el registro de paginas de Power BI con pregunta, tabla usada, medidas y estado.",
            "Documentar medidas DAX relevantes y su razon de uso.",
            "Insertar capturas del modelo Power BI y de cada visual validado.",
            "Validar nulos, duplicados y cobertura antes de interpretar hallazgos.",
            "Ampliar el analisis prescriptivo con criterios revisables y no automaticos.",
        ],
    )

    doc.add_heading("16. Anexos", level=1)
    doc.add_heading("16.1 Diccionario operativo minimo", level=2)
    add_kv_table(
        doc,
        [
            ("Carta", "Entidad base identificada por cards.card_id."),
            ("Aparicion", "Fila de card_sets que combina carta, set/codigo y rareza."),
            (
                "Marketplace",
                "Fuente de precio: Cardmarket, TCGPlayer, eBay, Amazon o CoolStuffInc.",
            ),
            (
                "Snapshot",
                "Foto historica de precios insertada por una ejecucion real del ETL.",
            ),
            (
                "Outlier",
                "Precio candidato a revision; no implica oportunidad ni error por si mismo.",
            ),
        ],
    )

    doc.add_heading("16.2 Registro de medidas DAX", level=2)
    add_matrix(
        doc,
        ["Medida", "Formula o descripcion", "Estado"],
        [["Pendiente", "Pendiente de definicion DAX", "Pendiente"]],
        [2500, 4860, 2000],
    )

    doc.add_heading("16.3 Registro de mantenimiento del informe", level=2)
    add_matrix(
        doc,
        ["Elemento", "Responsable", "Estado"],
        [["Modelo relacional Power BI", "Proyecto", "Vigente"]],
        [1800, 1800, 5760],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    try:
        build_doc()
        print(OUT_PATH)
    except PermissionError:
        OUT_PATH = BASE_DIR / "informe_analisis_powerbi_yugioh_actualizado.docx"
        build_doc()
        print(OUT_PATH)
